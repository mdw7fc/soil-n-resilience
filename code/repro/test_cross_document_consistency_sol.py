#!/usr/bin/env python3
"""Verify final manuscript/SI/response values against regenerated SOL outputs."""
from __future__ import annotations

import hashlib
from pathlib import Path
import zipfile

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
# The v17 release documents live in the tree, so this test runs from any
# checkout (F-026); it reads the CLEAN copies because python-docx does not
# surface text inside <w:ins> runs, so a tracked file under-reports its own
# accepted content.
SUBMISSION = ROOT / "resumbission" / "v17"

MANUSCRIPT = SUBMISSION / "Wallenstein-Manning_ERFS_manuscript_v17-clean.docx"
SI = SUBMISSION / "Wallenstein-Manning_ERFS_SI_v17-clean.docx"
RESPONSE = SUBMISSION / "Author_Response_ERFS-100341_v17-clean.docx"


def all_text(path: Path) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def require(text: str, fragments: list[str], label: str) -> None:
    missing = [item for item in fragments if item not in text]
    assert not missing, f"{label}: missing {missing}"


def forbid(text: str, fragments: list[str], label: str) -> None:
    present = [item for item in fragments if item in text]
    assert not present, f"{label}: stale fragments {present}"


def embedded_sha(path: Path, internal: str) -> str:
    with zipfile.ZipFile(path) as archive:
        return hashlib.sha256(archive.read(internal)).hexdigest()


def file_sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def main() -> None:
    manuscript = all_text(MANUSCRIPT)
    si = all_text(SI)
    response = all_text(RESPONSE)

    require(manuscript, [
        "3.0% by year 10",
        "5.1% (FSU/Central Asia)",
        "0.3–1.0 percentage points",
        "approximately 19%",
        "one-fifth of global seaborne fertilizer trade",
        "3.5 mm",
        "reduces Sub-Saharan African year-10 yield loss by approximately 8%",
        "ε_F,N = −0.50",
        "0.9–2.2 times the Century estimates",   # four-pool reinstated (F-029)
        "exploratory structural sensitivity",    # F-030 framing
        "larger in seven of eight regions",
        "alternative representations of SOM stabilization",  # F-031 scope
        "land-response coefficient",           # replaced gross-margin descriptor
    ], "manuscript")
    require(si, [
        "6.198", "6.022", "6.220", "3.773", "4.874", "5.414", "3.967", "4.318",
        "ρ = +0.19",
        "2.6% (East Asia) to 11.9%",
        "in none of the 1,000 draws",
        "default −0.50",
        "31.9",                      # derived BNF, Table 1 and prose
        "y_max = 3.97",
        "The median ratio is 1.65×",  # four-pool reinstated (F-029)
        "roughly half of MAOM formation",         # corrected pathway (F-030)
        "from 1.77 to 1.56",                      # texture sensitivity
        "exploratory structural sensitivity rather than a controlled experiment",
        "enumerated in the deposited parameter ledger",
        # F-031: mechanism attribution stays within what was demonstrated
        "cannot be uniquely attributed to a single mechanism",
        "combined effects of its alternative stabilization",
        "from 0.9× to 2.2× across regions",
        "equilibrium yield 1.37 t ha",   # SSA calibration shortfall disclosed
        "Model verification and reproducibility",   # Note 8, standalone
        "Crop calendar. Planting and maturity months are fixed",
        "Sacks W J",
        "ρ = +0.54",                 # fert-elasticity yr10, Note 3
        "plus 12,000 evaluations",   # Note 6: yield vs partial-net-revenue counts
        "in each of the four regions with audited price pairs",
    ], "SI")
    require(response, [
        "εF,N = −0.50; active in S3, SC1, and SC2",
        "~50%",
        "0.3–1.0 percentage points",
        "ρ = +0.19",
        "0.62 percentage points",
        "y_max = 3.97",
        "re-implemented the scheme on the deposited model",  # F-029
        "The relevant revised manuscript text now reads:",
        "Added at the end of the Introduction.",
        "evaluated against independent temperate and tropical benchmarks",
    ], "response")

    stale = [
        "6.277", "3.876", "6.119", "6.095",
        "ρ = +0.02", "ρ = +0.07", "ρ = +0.60", "ρ = −0.86", "ρ = −0.45",
        "yield loss by approximately 55%",
        "one-third of global fertilizer trade",
        "2.6–4.6 percentage points",
        "~25%",
        "8.4 mm WHC",
        "y_max = 3.88",
        "0.54 percentage points",
        "roughly ten times",
        "roughly twice the year-30",
        "83.7% of draws despite",
        "2.6 to 10.3 percentage points",
        # F-028: four-pool withdrawal must hold everywhere, and the partial
        # net-revenue framing admits no leftover gross-margin claims
        "Gross-margin losses follow",
        "fertilizer share in regional gross margin",
        "gross-margin change",
        "higher year-1 gross margin",
        "in gross-margin terms",
        "gross-margin response",
        "retained here as originally reported",
        "we retain this analysis",
        "comparable or larger under halved",
        "3.0–5.9 percentage points when halved",
    ]

    # Version-referential language is banned from the two PUBLISHED documents
    # (only one version is ever published; revision history belongs in the
    # response letter alone), together with the superseded four-pool family
    # and the stale figS10 caption values (F-029).
    stale_published = [
        # F-030: the four-pool description and interpretation must stay
        # inside what the code and the baselines support
        "all SOM formation",
        "the ratio is unaffected",
        "while amplifying its magnitude",
        "the only region below parity",
        "conservative choice",
        "must pass through microbial assimilation before entering",
        # F-031: the fixed-CUE analysis partitions only 6-59% of the loss
        # difference; single-mechanism causal attributions are unsupported
        "The dominant mechanism is",
        "dominated by respiratory losses",
        "because carbon routed through microbial biomass",
        "microbially driven stabilization feedbacks",
        "submitted version",
        "in this revision",
        "added in revision",
        "0.9× to 2.5×",
        "median ratio is 1.8×",
        "2.5× in Latin America",
        "21.0% to 9.4%",
        "matched 20% SC1",
        "Model internal-consistency corrections",
        "and and,",
    ]
    forbid(manuscript + "\n" + si, stale_published, "manuscript+SI")
    forbid(manuscript + "\n" + si + "\n" + response, stale, "all documents")

    table = pd.read_csv(ROOT / "outputs" / "Table_S4_calibration_sol.csv")
    table = table.drop_duplicates(subset="region", keep="first").reset_index(drop=True)
    si_doc = Document(SI)
    table_s4 = next(
        t for t in si_doc.tables
        if len(t.rows) > 0 and len(t.rows[0].cells) >= 8
        and t.rows[0].cells[0].text.strip() == "Region"
        and "y_max" in t.rows[0].cells[4].text
    )
    labels = [
        "N America", "Europe", "E Asia", "S Asia", "SE Asia",
        "L America", "Sub-Saharan Africa", "FSU/C Asia",
    ]
    for label, row, record in zip(
        labels, table_s4.rows[1:], table.itertuples(index=False)
    ):
        assert row.cells[0].text.strip() == label
        assert abs(float(row.cells[4].text) - record.calibrated_y_max_t_ha) <= .01
        assert abs(
            float(row.cells[7].text)
            - record.simulated_year2_no_synth_n_t_ha
        ) <= .01

    media = [
        (MANUSCRIPT, "word/media/image1.png", ROOT / "figures/Figure_1_farm_buffering.png"),
        (MANUSCRIPT, "word/media/image2.png", ROOT / "figures/Figure_2_regional_vulnerability.png"),
        (SI, "word/media/image4.png", ROOT / "figures/Figure_S4_hindcast_sensitivity.png"),
        (SI, "word/media/image5.png", ROOT / "figures/Figure_S5_fourpool_flux.png"),
        (SI, "word/media/image6.png", ROOT / "figures/Figure_S6_pairwise_diagnostics.png"),
        (SI, "word/media/image8.png", ROOT / "figures/Figure_S8_elasticity_sensitivity.png"),
        (SI, "word/media/image10.png", ROOT / "figures/Figure_S10_nue_sensitivity.png"),
        (SI, "word/media/image12.png", ROOT / "figures/Figure_S12_crop_response_calibration.png"),
        (SI, "word/media/image13.png", ROOT / "figures/Figure_S13_OFRA_SSA_validation.png"),
        (SI, "word/media/image9.png", ROOT / "figures/Figure_S9_mc_ensemble.png"),
        (MANUSCRIPT, "word/media/image3.png", ROOT / "figures/Figure_3_mechanism_screen.png"),
    ]
    for docx, internal, source in media:
        assert embedded_sha(docx, internal) == file_sha(source), (
            f"{docx.name}:{internal} does not match {source.name}"
        )

    print("CROSS-DOCUMENT CONSISTENCY: PASS")
    print("  headline values, BNF/buffer tables, Table S4, benchmark language")
    print("  and embedded regenerated figures agree with the SOL outputs.")


if __name__ == "__main__":
    main()
